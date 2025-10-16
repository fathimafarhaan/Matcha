import os
import torch
import re
import numpy as np
import joblib
import spacy
import json
import jwt as pyjwt
import bcrypt
import csv
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pdfminer.high_level import extract_text
from docx import Document
from werkzeug.utils import secure_filename
from transformers import DistilBertTokenizer, DistilBertModel
from datetime import datetime, timedelta
from functools import wraps
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.shared import qn
from docx.enum.section import WD_SECTION


app = Flask(__name__)
CORS(app)

# Directory for protected CSV storage
PROTECTED_DIR = "protected_data"
os.makedirs(PROTECTED_DIR, exist_ok=True)
CSV_FILE_PATH = os.path.join(PROTECTED_DIR, "resumes.csv")

SECRET_KEY = "secret_key"  # Should be loaded from environment variables
UPLOAD_FOLDER = "uploads"
SCORING_CONFIG_FILE = "scoring_config.json"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Sample users (hashed passwords for security)
users = {
    "recruiter": {
        "username": "recruiter",
        "password": bcrypt.hashpw("recruiter123".encode(), bcrypt.gensalt()).decode(),
        "role": "recruiter"
    },
    "admin": {
        "username": "admin",
        "password": bcrypt.hashpw("admin123".encode(), bcrypt.gensalt()).decode(),
        "role": "admin"
    }
}

# Lazy loading for ML models
_model = None
_encoder = None
_scaler = None
_tokenizer = None
_bert_model = None
_nlp = None

def get_model():
    global _model
    if _model is None:
        _model = joblib.load("backend/resume_model.pkl")
    return _model

def get_encoder():
    global _encoder
    if _encoder is None:
        _encoder = joblib.load("backend/encoder.pkl")
    return _encoder

def get_scaler():
    global _scaler
    if _scaler is None:
        _scaler = joblib.load("backend/scaler.pkl")
    return _scaler

def get_tokenizer():
    global _tokenizer
    if _tokenizer is None:
        _tokenizer = DistilBertTokenizer.from_pretrained("distilbert-base-uncased")
    return _tokenizer

def get_bert_model():
    global _bert_model
    if _bert_model is None:
        _bert_model = DistilBertModel.from_pretrained("distilbert-base-uncased")
    return _bert_model

def get_nlp():
    global _nlp
    if _nlp is None:
        _nlp = spacy.load("en_core_web_sm")
    return _nlp

# Load scoring configuration
def load_scoring_config():
    if not os.path.exists(SCORING_CONFIG_FILE):
        return {"experience_weight": 0.3, "projects_weight": 0.2, "skills_weight": 0.3, "education_weight": 0.2}
    with open(SCORING_CONFIG_FILE, "r") as file:
        return json.load(file)

scoring_config = load_scoring_config()

# Authentication decorator
def token_required(role=None):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            token = request.headers.get("Authorization")
            if not token or not token.startswith("Bearer "):
                return jsonify({"error": "Unauthorized. Token required."}), 403
            
            user = verify_token(token.split("Bearer ")[1])
            if not user:
                return jsonify({"error": "Invalid or expired token."}), 403
                
            if role and user.get("role") != role:
                return jsonify({"error": f"Forbidden. {role.capitalize()} access only."}), 403
                
            return f(user, *args, **kwargs)
        return decorated
    return decorator

# Generate JWT token
def generate_token(username, role):
    payload = {"username": username, "role": role, "exp": datetime.utcnow() + timedelta(hours=2)}
    return pyjwt.encode(payload, SECRET_KEY, algorithm="HS256")

# JWT Verification
def verify_token(token):
    try:
        return pyjwt.decode(token, SECRET_KEY, algorithms=["HS256"])
    except (pyjwt.ExpiredSignatureError, pyjwt.InvalidTokenError):
        return None

# Login Route
@app.route("/login", methods=["POST"])
def login():
    data = request.json
    username, password = data.get("username"), data.get("password")
    if username in users and bcrypt.checkpw(password.encode(), users[username]["password"].encode()):
        return jsonify({"token": generate_token(username, users[username]["role"]), "role": users[username]["role"]})
    return jsonify({"error": "Invalid username or password"}), 401

# Verify Token Route
@app.route("/verify-token", methods=["POST"])
def verify_token_endpoint():
    token = request.headers.get("Authorization")
    if not token or not token.startswith("Bearer "):
        return jsonify({"error": "Token required"}), 401
    
    user = verify_token(token.split("Bearer ")[1])
    if not user:
        return jsonify({"error": "Invalid or expired token"}), 401
    
    return jsonify({"valid": True, "user": user}), 200

#extract from .docx
def extract_docx_text_with_headers(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        # Extract text from main body
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
        
        # Extract header and footer text
        for section in doc.sections:
            # Header text extraction
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Check for tables in header
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    full_text.append(para.text)
            
            # Footer text extraction
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Check for tables in footer
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    full_text.append(para.text)
        
        return "\n".join(full_text) if full_text else None
    except PackageNotFoundError:
        print("Package not found error when opening the document")
        return None
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return None

# Resume Text Extraction
def extract_resume_text(file):
    try:
        file_path = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
        file.save(file_path)
        
        if file.filename.endswith(".pdf"):
            text = extract_text(file_path)
        elif file.filename.endswith((".docx", ".doc")):
            text = extract_docx_text_with_headers(file_path)
        else:
            os.remove(file_path)
            return None
            
        os.remove(file_path)
        return text.strip() if text.strip() else None
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        print(f"Error extracting text: {str(e)}")
        return None

# DistilBERT Text Embedding
def get_text_embedding(text, chunk_size=512):
    try:
        tokenizer = get_tokenizer()
        bert_model = get_bert_model()
        
        tokens = tokenizer(text, return_tensors="pt", truncation=False, padding=False)
        input_ids, embeddings = tokens["input_ids"].squeeze(), []
        
        for i in range(0, len(input_ids), chunk_size):
            with torch.no_grad():
                chunk = input_ids[i: i + chunk_size].unsqueeze(0)
                embeddings.append(bert_model(input_ids=chunk).last_hidden_state.mean(dim=1).numpy())
                
        return np.mean(embeddings, axis=0)
    except Exception as e:
        print(f"Error in text embedding: {str(e)}")
        # Return zero vector as fallback
        return np.zeros((768,))

# Extract name, email and phone
def extract_contact_info(text):
    # Email extraction
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    email_match = re.search(email_pattern, text)
    email = email_match.group(0) if email_match else "N/A"
    
    # Phone extraction  
    phone_pattern = r"\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}"
    phone_match = re.search(phone_pattern, text)
    phone = phone_match.group(0) if phone_match and len(re.sub(r"\D", "", phone_match.group(0))) >= 10 else "N/A"
    
    # Name extraction with improved methods
    name = extract_name(text)
    
    return name, email, phone

def extract_name(text):
    nlp = get_nlp()
    
    # First attempt: Look for patterns like "Name: John Smith" or "John Smith - Resume"
    name_patterns = [
        r"(?i)^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?:\n|$)",  # Name at the start of resume (on same line)
        r"(?i)name:?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?:\n|$)",  # "Name: John Smith"
        r"(?i)([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})'s\s+resume(?:\n|$)",  # "John Smith's Resume"
        r"(?i)([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\s+-\s+resume(?:\n|$)",  # "John Smith - Resume"
        r"(?i)resume\s+of\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?:\n|$)",  # "Resume of John Smith"
        r"(?i)curriculum\s+vitae\s+of\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})(?:\n|$)",  # "CV of John Smith"
    ]
    
    for pattern in name_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    # Second attempt: Check first 200 characters for capitalized words within single lines
    first_section = text[:200]
    lines = [line.strip() for line in first_section.split('\n') if line.strip()]
    
    for line in lines:
        # Skip lines with common header words
        if any(word in line.lower() for word in ["resume", "cv", "curriculum", "vitae", "contact", "address", "phone", "email"]):
            continue
            
        # Check if line contains a potential name (2-4 capitalized words)
        words = line.split()
        if len(words) >= 2 and len(words) <= 4:  # Names are typically 2-4 words
            if all(word[0].isupper() for word in words if len(word) > 1):
                return line.strip()
    
    # Third attempt: Use spaCy's NER for PERSON entities, but only within single lines
    doc = nlp(text[:500])  # Process only first 500 chars for efficiency
    
    for sent in doc.sents:
        person_entities = [ent.text for ent in sent.ents if ent.label_ == "PERSON"]
        if person_entities:
            return person_entities[0]
    
    # Fallback to first capitalized name pattern within a single line
    for line in text[:500].split('\n'):
        name_match = re.findall(r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})", line)
        if name_match:
            return name_match[0]
    
    return "Unknown"
# Combined entity extraction function
def extract_resume_entities(text):
    nlp = get_nlp()
    doc = nlp(text)
    
    # Initialize collections
    skills = set()
    education = set()
    certifications = set()
    job_role = set()
    
    # For storing experience in original units
    experience_data = {
        "years": 0,
        "months": 0,
        "weeks": 0,
        "days": 0,
        "original_unit": "months",  # Default return unit
        "original_value": 0
    }
    
    # Skills extraction - keyword based
    skill_keywords = ["python", "java", "mysql","c++", "tensorflow", "pytorch", "deep learning", 
                     "machine learning", "r","react","node.js","html","css", "javascript", 
                     "docker", "nextjs", "nodejs"]
    for skill in skill_keywords:
        if re.search(r'\b' + re.escape(skill) + r'\b', text.lower()):
            skills.add(skill)
    
    # Job role extraction based on job title keywords
    job_title_keywords = [
        "data scientist", "software engineer", "backend developer", "frontend developer", 
        "full stack developer", "ml engineer", "machine learning engineer", "data analyst",
        "product manager", "project manager", "devops engineer", "cloud engineer",
        "web developer", "mobile developer", "qa engineer", "system administrator",
        "ux designer", "ui developer", "database administrator", "business analyst",
        "software developer", "software developer intern", "intern", "ai/ml engineer",
        "frontend developer", "fullstack developer"
    ]
    
    # Extract job roles using regex pattern matching
    for title in job_title_keywords:
        if re.search(r"\b" + re.escape(title) + r"\b", text.lower()):
            job_role.add(title.title())
    
    # Look for explicit work experience sections
    work_exp_section = None
    work_section_headers = ["work experience", "employment", "professional experience"]
    text_lower = text.lower()
    
    for header in work_section_headers:
        if header in text_lower:
            start_idx = text_lower.find(header)
            section_text = text[start_idx:start_idx + 1000]  # Look at 1000 chars after section header
            work_exp_section = section_text
            break
    
    # Process explicit work experience
    if work_exp_section:
        time_patterns = [
            (r"(\d+\.?\d*)\s*(?:years?|yrs?)", "years"),
            (r"(\d+\.?\d*)\s*(?:months?|month)", "months"),
            (r"(\d+\.?\d*)\s*(?:weeks?|week)", "weeks"),
            (r"(\d+\.?\d*)\s*(?:days?|day)", "days")
        ]
        
        for pattern, unit in time_patterns:
            matches = re.findall(pattern, work_exp_section, re.IGNORECASE)
            for match in matches:
                value = float(match)
                experience_data[unit] += value
                
                # Remember the original unit and value for the first match
                if experience_data["original_value"] == 0:
                    experience_data["original_unit"] = unit
                    experience_data["original_value"] = value
    
    
        for pattern, unit in time_patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                value = float(match)
                experience_data[unit] += value
                
                # Remember the original unit and value for the first match
                if experience_data["original_value"] == 0:
                    experience_data["original_unit"] = unit
                    experience_data["original_value"] = value
    
    # For backward compatibility, still calculate years for the model
    experience = (
        experience_data["years"] + 
        experience_data["months"] / 12.0 + 
        experience_data["weeks"] / 52.0 + 
        experience_data["days"] / 365.0
    )
    
    # Create the experience string for display
    experience_display = None
    if experience_data["original_value"] > 0:
        experience_display = f"{experience_data['original_value']} {experience_data['original_unit']}"
    
    # Certification extraction using specific keywords and patterns
    cert_patterns = [
        r"(?i)certified\s+(\w+(?:\s+\w+){0,5})",
        r"(?i)(\w+(?:\s+\w+){0,3})\s+certification",
        r"(?i)certificate\s+in\s+(\w+(?:\s+\w+){0,5})",
        r"(?i)completed\s+(?:course|training|program)(?:\s+in)?\s+(\w+(?:\s+\w+){0,5})"
    ]
    
    for pattern in cert_patterns:
        cert_matches = re.finditer(pattern, text)
        for match in cert_matches:
            certification = match.group(1).strip()
            if len(certification) > 3:  # Filter out very short matches
                certifications.add(certification)
    
    #NER
    for ent in doc.ents:
        if ent.label_ in ["WORK_OF_ART", "ORG", "PRODUCT", "EVENT"]:
            if "certification" in ent.sent.text.lower() or "certified" in ent.sent.text.lower():
                cert_text = ent.text.strip()
                if len(cert_text) > 3:
                    certifications.add(cert_text)
    
    # Extract education details
    education_details = extract_education_details(text)
    education.update(education_details)
    
    # Clean and filter education results
    education = clean_education_results(list(education))
    
    # Extract projects
    projects, projects_count = extract_projects(text)
    
    # Return both the calculated years (for model) and the display string
    return list(skills), education, list(certifications), list(job_role), experience, projects, projects_count, experience_display

def extract_education_details(text):
    """Extract education details using multiple techniques"""
    nlp = get_nlp()
    education = set()
    
    # Define education keywords and degree patterns
    edu_indicators = [
        "university", "college", "institute", "school", "academy",
        "bachelor", "master", "ph.d", "phd", "doctorate", "diploma", "degree",
        "b.tech", "m.tech", "mba", "b.sc", "m.sc", "b.a", "m.a",
        "engineering", "business administration", "computer science"
    ]
    
    degree_patterns = [
        r"(?i)(?:bachelor|master|doctorate|ph\.?d\.?|b\.?tech|m\.?tech|mba|b\.?sc|m\.?sc|b\.?a|m\.?a)[\s.]+((?:of|in)\s+)?([a-zA-Z\s]+(?:science|engineering|technology|administration|management|arts|commerce))",
        r"(?i)(bachelor|master|doctorate|ph\.?d\.?|b\.?tech|m\.?tech|mba|b\.?sc|m\.?sc|b\.?a|m\.?a)[\s.]+(?:degree)?",
        r"(?i)(university|college|institute|school)(?:\s+of)?\s+([a-zA-Z\s&]+)"
    ]
    
    # Find education sections
    education_section = None
    text_lower = text.lower()
    edu_section_headers = ["education", "academic background", "qualifications", "academic", "educational"]
    
    for header in edu_section_headers:
        if header in text_lower:
            # Extract section following the header until the next header
            start_idx = text_lower.find(header)
            section_text = text[start_idx:start_idx + 1000]  # Look at 1000 chars after section header
            education_section = section_text
            break
    
    # First approach: Use spaCy to find ORG entities within the education section
    if education_section:
        doc = nlp(education_section)
        for ent in doc.ents:
            if ent.label_ == "ORG":
                # Verify if it's likely an educational institution
                if any(indicator in ent.text.lower() for indicator in edu_indicators):
                    education.add(ent.text)
    
    
    return education

def clean_education_results(education_list):
    """Clean up and filter education results to remove duplicates and irrelevant items"""
    if not education_list:
        return []
        
    # Remove items that are too short (likely noise)
    filtered = [edu for edu in education_list if len(edu) > 5]
    
    # Remove items that are too generic
    blacklist = ["company", "corporation", "inc", "llc", "resume", "cv"]
    filtered = [edu for edu in filtered if not any(term in edu.lower() for term in blacklist)]
    
    # Remove duplicates while preserving order
    unique_edu = []
    seen = set()
    for edu in filtered:
        if edu.lower() not in seen:
            seen.add(edu.lower())
            unique_edu.append(edu)
    
    return unique_edu

def extract_projects(text):
    """
    Extract projects from resume text with improved accuracy.
    Returns a list of project details and the total count.
    """
    projects = []
    text_lower = text.lower()
    
    # Find projects section
    project_section = None
    project_section_headers = ["projects", "project experience", "personal projects", 
                             "technical projects", "academic projects", "professional projects"]
    
    # Try to locate the projects section
    for header in project_section_headers:
        pattern = r"(?i)(?:^|\n|\r)[\s]*" + re.escape(header) + r"[\s]*(?::|\n|\r|$)"
        match = re.search(pattern, text_lower)
        if match:
            start_idx = match.start()
            # Find the next section header or end of text
            next_section_match = re.search(r"(?i)(?:^|\n|\r)[\s]*(?:experience|education|skills|certifications|achievements|languages|references)[\s]*(?::|\n|\r|$)", text_lower[start_idx + len(header):])
            end_idx = start_idx + len(header) + next_section_match.start() if next_section_match else len(text)
            project_section = text[start_idx:end_idx]
            break
    
    # If we found a projects section, parse it
    if project_section:
        # Look for project titles/entries in the section
        project_entries = re.findall(r"(?i)(?:^|\n|\r)[\s]*(?:•|\-|\*|\d+\.|\(\d+\)|[A-Z][a-zA-Z\s]+:)[\s]*(.*?)(?=(?:^|\n|\r)[\s]*(?:•|\-|\*|\d+\.|\(\d+\)|[A-Z][a-zA-Z\s]+:)|$)", project_section, re.MULTILINE | re.DOTALL)
        for entry in project_entries:
            entry = entry.strip()
            if entry and len(entry) > 10:  # Avoid very short entries that might be noise
                projects.append(entry)
    
    # If no projects section or no entries found, try alternative approach
    if not projects:
        # Look for project keyword patterns throughout the document
        project_patterns = [
            r"(?i)(?:^|\n|\r)[\s]*(?:•|\-|\*|\d+\.|\(\d+\))[\s]*(?:[A-Z][a-zA-Z\s]+[\s]*[-–][\s]*.*?project.*?)(?=(?:^|\n|\r)[\s]*(?:•|\-|\*|\d+\.|\(\d+\))|$)",
            r"(?i)(?:developed|created|implemented|designed|built)[\s]*(?:a|an)?[\s]*(?:.*?)[\s]*(?:project|application|system|website|software|tool|platform)",
            r"(?i)project[\s]*:[\s]*([^\n\r]+)"
        ]
        
        for pattern in project_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.DOTALL)
            for match in matches:
                entry = match.group(0).strip()
                if entry and len(entry) > 10 and not any(proj in entry for proj in projects):
                    projects.append(entry)
    
    # Count GitHub/repository mentions as potential additional projects
    github_repos = re.findall(r"(?i)github\.com/[a-zA-Z0-9_-]+/([a-zA-Z0-9_-]+)", text)
    gitlab_repos = re.findall(r"(?i)gitlab\.com/[a-zA-Z0-9_-]+/([a-zA-Z0-9_-]+)", text)
    
    # Add repositories that aren't already in projects list
    for repo in github_repos + gitlab_repos:
        repo_name = repo.replace("-", " ").replace("_", " ")
        if not any(repo_name.lower() in proj.lower() for proj in projects):
            projects.append(f"Repository: {repo_name}")
    
    # Final verification to avoid counting non-projects
    verified_projects = []
    for proj in projects:
        # Skip entries that are likely not projects
        if any(keyword in proj.lower() for keyword in ["not applicable", "n/a", "none"]):
            continue
        verified_projects.append(proj)
    
    return verified_projects, len(verified_projects)

def save_resume_to_csv(data):
    """Save resume data to CSV file"""
    try:
        file_exists = os.path.isfile(CSV_FILE_PATH)
        
        # Make sure the directory exists
        os.makedirs(os.path.dirname(CSV_FILE_PATH), exist_ok=True)
        
        with open(CSV_FILE_PATH, mode="a", newline="", encoding="utf-8") as file:
            writer = csv.DictWriter(file, fieldnames=data.keys())
            if not file_exists:
                writer.writeheader()
            writer.writerow(data)
        return True
    except Exception as e:
        print(f"Error saving to CSV: {str(e)}")
        return False

# Resume Upload Route
@app.route("/upload", methods=["POST"])
@token_required(role="recruiter")
def upload_resume(user):
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files["file"]
    
    # Extract text from resume
    if not (resume_text := extract_resume_text(file)):
        return jsonify({"error": "Failed to extract text"}), 400

    # Get text embedding for ML model
    embedding = get_text_embedding(resume_text)
    
    # Extract all entities from resume
    skills, education, certifications, job_role, experience, projects_list, projects_count, experience_display = extract_resume_entities(resume_text)
    # Get contact information
    name, email, phone = extract_contact_info(resume_text)
    
    # Prepare data for ML model
    encoder = get_encoder()
    scaler = get_scaler()
    model = get_model()
    
    # Create a simple validation for the encoder
    validated_data = [
        skills[0] if skills else "Unknown",
        education[0] if education else "Unknown", 
        certifications[0] if certifications else "Unknown",
        job_role[0] if job_role else "Unknown"
    ]
    
    # For now, create dummy encoded data since our encoder is simplified
    # The model expects 10 features, so we'll create dummy data
    encoded_cats = np.array([[0, 0, 0, 0, 0, 0, 0, 0]])  # 8 dummy encoded features
    scaled_nums = scaler.transform([[experience, projects_count]])
    final_features = np.hstack((encoded_cats, scaled_nums))  # Total: 8 + 2 = 10 features

    # Calculate base score from ML model
    base_score = model.predict_proba(final_features)[0][1] * 100

    # Apply scoring weights
    weighted_score = (
        (experience * scoring_config["experience_weight"]) +
        (projects_count * scoring_config["projects_weight"]) +
        (len(skills) * scoring_config["skills_weight"]) +
        (len(education) * scoring_config["education_weight"])
    )
    
    # Calculate final score
    final_score = round((base_score + weighted_score) / 2, 2)

    return jsonify({
        "parsedResume": resume_text[:1000], 
        "score": final_score,
        "name": name,
        "phone": phone,
        "email": email,
        "experience": experience,  # Keep the numerical value for calculations
        "experience_display": experience_display,  # Add the display string
        "skills": skills,
        "education": education,
        "projects": projects_count
    })

@app.route("/save_csv", methods=["POST"])
@token_required(role="recruiter")
def save_csv(user):
    data = request.json
    required_fields = ["name", "phone", "email", "experience", "skills", "education", "projects", "score"]
    
    # Validate required fields
    if not all(field in data for field in required_fields):
        return jsonify({"error": "Missing required fields"}), 400
    
    # Convert skills and education to comma-separated strings if they're lists
    if isinstance(data.get("skills"), list):
        data["skills"] = ", ".join(data["skills"])
    if isinstance(data.get("education"), list):
        data["education"] = ", ".join(data["education"])
    
    try:
        if save_resume_to_csv(data):
            return jsonify({"message": "Resume details saved to CSV successfully"}), 200
        else:
            return jsonify({"error": "Failed to save resume details"}), 500
    except Exception as e:
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/admin/download_csv", methods=["GET"])
@token_required(role="admin")
def download_csv_admin(user):
    if not os.path.exists(CSV_FILE_PATH):
        return jsonify({"error": "CSV file not found"}), 404

    return send_file(CSV_FILE_PATH, as_attachment=True)

@app.route("/admin/get_scoring_criteria", methods=["GET"])
@token_required(role="admin")
def get_scoring_criteria(user):
    return jsonify(scoring_config), 200

@app.route("/admin/update_scoring", methods=["POST"])
@token_required(role="admin")
def update_scoring(user):
    data = request.json
    required_keys = {"experience_weight", "projects_weight", "skills_weight", "education_weight"}
    if not all(k in data for k in required_keys):
        return jsonify({"error": "Invalid scoring criteria"}), 400

    try:
        with open(SCORING_CONFIG_FILE, "w") as file:
            json.dump(data, file, indent=4)

        global scoring_config
        scoring_config = load_scoring_config()
        return jsonify({"message": "Scoring criteria updated successfully"}), 200
    except Exception as e:
        return jsonify({"error": f"Failed to update scoring criteria: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True)
